import streamlit as st
import re
from datetime import datetime
from docx import Document
from io import BytesIO
import os

st.set_page_config(page_title="📅 Date Changer", layout="centered")
st.title("📁 Date Replacer for .docx and .txt Files")

# Build regex pattern for flexible date detection
def build_flexible_date_regex():
    months = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|" \
             r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|" \
             r"Nov(?:ember)?|Dec(?:ember)?)"
    patterns = [
        r"\d{1,2}/\d{1,2}/\d{4}",      # dd/mm/yyyy or mm/dd/yyyy
        r"\d{1,2}-\d{1,2}-\d{4}",      # dd-mm-yyyy or mm-dd-yyyy
        r"\d{4}/\d{1,2}/\d{1,2}",      # yyyy/mm/dd
        r"\d{4}-\d{1,2}-\d{1,2}",      # yyyy-mm-dd
        r"\d{4}\.\d{1,2}\.\d{1,2}",  # yyyy.mm.dd
        r"\d{1,2}\.\d{1,2}\.\d{4}",  # dd.mm.yyyy
        rf"\d{{1,2}} ?{months} ?\d{{4}}",
        rf"\d{{1,2}}{months}\d{{4}}",
        rf"{months}/\d{{1,2}}",
        rf"{months}\d{{2}}",
        rf"\d{{2}}{months}"
    ]
    return "|".join(patterns)

# Replace and preserve date format
def format_preserving_replace(match, new_date_obj):
    old = match.group(0)
    day = new_date_obj.strftime('%d')
    month = new_date_obj.strftime('%B')
    mon_abbr = new_date_obj.strftime('%b')
    mon_number = new_date_obj.strftime('%m')
    year = new_date_obj.strftime('%Y')

    if re.match(r"\d{1,2}/\d{1,2}/\d{4}", old):
        return f"{day}/{mon_number}/{year}"
    elif re.match(r"\d{1,2}-\d{1,2}-\d{4}", old):
        return f"{day}-{mon_number}-{year}"
    elif re.match(r"\d{4}/\d{1,2}/\d{1,2}", old):
        return f"{year}/{mon_number}/{day}"
    elif re.match(r"\d{4}-\d{1,2}-\d{1,2}", old):
        return f"{year}-{mon_number}-{day}"
    elif re.match(r"\d{4}\.\d{1,2}\.\d{1,2}", old):
        return f"{year}.{mon_number}.{day}"
    elif re.match(r"\d{1,2}\.\d{1,2}\.\d{4}", old):
        return f"{day}.{mon_number}.{year}"
    elif re.match(r"\d{1,2}(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\d{4}", old, re.IGNORECASE):
        return f"{day}{mon_abbr}{year}"
    elif re.match(r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\d{2}", old, re.IGNORECASE):
        return f"{mon_abbr}{day}"
    elif re.match(r"(January|February|March|April|May|June|July|August|September|October|November|December)/\d{1,2}", old, re.IGNORECASE):
        return f"{month}/{day}"
    elif re.match(r"\d{1,2}(January|February|March|April|May|June|July|August|September|October|November|December)\d{4}", old, re.IGNORECASE):
        return f"{day}{month}{year}"
    elif re.match(r"\d{1,2} ?(January|February|March|April|May|June|July|August|September|October|November|December) ?\d{4}", old, re.IGNORECASE):
        return f"{day} {month} {year}"
    else:
        return new_date_obj.strftime('%Y-%m-%d')

# Replace in plain text
def replace_dates_in_text(content, new_date_obj, pattern):
    return pattern.sub(lambda m: format_preserving_replace(m, new_date_obj), content), len(pattern.findall(content))

# Replace in docx paragraph runs
def replace_in_runs(paragraph, new_date_obj, date_pattern):
    combined_text = ''.join(run.text for run in paragraph.runs)
    matches = list(date_pattern.finditer(combined_text))
    if not matches:
        return 0

    for match in reversed(matches):
        start, end = match.span()
        replacement = format_preserving_replace(match, new_date_obj)
        combined_text = combined_text[:start] + replacement + combined_text[end:]

    char_index = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run.text = combined_text[char_index:char_index + run_len]
        char_index += run_len

    return len(matches)

# Main processing function
def process_file_streamlit(uploaded_file, target_date, original_filename):
    ext = os.path.splitext(original_filename)[1].lower()
    filename_only = os.path.splitext(original_filename)[0]
    output_filename = f"{filename_only}_updated{ext}"
    pattern = re.compile(build_flexible_date_regex(), flags=re.IGNORECASE)
    new_date_obj = datetime.strptime(str(target_date), "%Y-%m-%d")

    buffer = BytesIO()

    if ext == ".txt":
        content = uploaded_file.read().decode("utf-8", errors="replace")
        updated_content, _ = replace_dates_in_text(content, new_date_obj, pattern)
        buffer.write(updated_content.encode("utf-8"))

    elif ext == ".docx":
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            replace_in_runs(para, new_date_obj, pattern)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_runs(para, new_date_obj, pattern)
        for section in doc.sections:
            for part in [section.header, section.footer]:
                for para in part.paragraphs:
                    replace_in_runs(para, new_date_obj, pattern)
        doc.save(buffer)

    else:
        raise ValueError("Unsupported file type")

    buffer.seek(0)
    return buffer, output_filename

# Streamlit interface
uploaded_file = st.file_uploader("Upload a .docx or .txt file", type=["docx", "txt"])
target_date = st.date_input("Choose the new target date")

if uploaded_file:
    filename = uploaded_file.name
    if st.button("🔁 Replace Dates"):
        try:
            buffer, new_filename = process_file_streamlit(uploaded_file, target_date, filename)
            st.success("✅ Dates replaced successfully!")
            st.download_button(
                label="📥 Download Modified File",
                data=buffer,
                file_name=new_filename,
                mime="application/octet-stream"
            )
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")


